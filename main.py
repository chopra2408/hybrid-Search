from rank_bm25 import BM25Okapi
from io import BytesIO
from langchain_groq import ChatGroq
from langchain_ollama import OllamaEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import ChatPromptTemplate
from langchain.schema import Document
from dotenv import load_dotenv
import pdfplumber
import json
import os
import faiss
import numpy as np
import streamlit as st
import docx

load_dotenv()
groq_api_key = os.getenv("GROQ_API_KEY")
cohere_api_key = os.getenv("COHERE_API_KEY")

st.title("Hybrid Search)")

if "llm" not in st.session_state:
    st.session_state.llm = ChatGroq(groq_api_key=groq_api_key, model_name="Llama3-8b-8192")

if "embeddings" not in st.session_state:
    st.session_state.embeddings = OllamaEmbeddings(model="nomic-embed-text")

if "faiss_index" not in st.session_state:
    st.session_state.faiss_index = None

if "documents_list" not in st.session_state:
    st.session_state.documents_list = []

if "bm25" not in st.session_state:
    st.session_state.bm25 = None

prompt = ChatPromptTemplate.from_template("""Answer the questions based on the provided context only. Provide the most accurate response.
You are allowed to use your own knowledge and be abusive too.
<context>
{context}
<context>
Questions: {input}
""")

def parse_files(files):
    documents = []
    for file in files:
        try:
            if file.type == "application/pdf":
                with pdfplumber.open(BytesIO(file.read())) as pdf:
                    content = "".join([page.extract_text() for page in pdf.pages if page.extract_text()])
                    if content:
                        documents.append(Document(page_content=content))
            elif file.type == "text/plain":
                content = file.read().decode("utf-8")
                if content:
                    documents.append(Document(page_content=content))
            elif file.type == "application/json":
                content = json.load(file)
                text_content = json.dumps(content)
                documents.append(Document(page_content=text_content))
            elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                doc = docx.Document(BytesIO(file.read()))
                content = "\n".join([para.text for para in doc.paragraphs])
                if content:
                    documents.append(Document(page_content=content))
            else:
                st.warning(f"Unsupported file format: {file.name}")
        except Exception as e:
            st.error(f"Error processing file {file.name}: {e}")
    return documents

def vector_embedding(files):
    with st.spinner("Processing and embedding documents..."):
        documents = parse_files(files)
        if not documents:
            st.error("No valid text extracted from the uploaded files.")
            return

        text_splitter = RecursiveCharacterTextSplitter(chunk_size=300, chunk_overlap=100)
        chunks = text_splitter.split_documents(documents)

        texts = []
        embeddings = []

        for chunk in chunks:
            embedding = st.session_state.embeddings.embed_query(chunk.page_content)
            texts.append(chunk.page_content)
            embeddings.append(embedding)

        embeddings_np = np.array(embeddings).astype("float32")

        if st.session_state.faiss_index is None:
            dimension = embeddings_np.shape[1]
            st.session_state.faiss_index = faiss.IndexFlatL2(dimension)

        st.session_state.faiss_index.add(embeddings_np)
        st.session_state.documents_list = texts

        st.session_state.bm25 = BM25Okapi([doc.split() for doc in texts])
        st.success("FAISS and BM25 hybrid search DB is ready.")

uploaded_files = st.file_uploader("Upload files (PDF, Text, JSON, DOCX, etc.)", type=["pdf", "txt", "json", "docx"], accept_multiple_files=True)

user_query = st.text_input("Enter your question based on the uploaded files:")

if st.button("Initialize Document Embeddings"):
    if uploaded_files:
        vector_embedding(uploaded_files)
    else:
        st.warning("Please upload files first.")

def hybrid_search(query, k=5):
    if not st.session_state.faiss_index or not st.session_state.documents_list:
        st.warning("FAISS index or documents are not initialized.")
        return []

    if st.session_state.bm25 is None:
        st.warning("BM25 is not initialized. Please upload documents first.")
        return []

    query_embedding = st.session_state.embeddings.embed_query(query)
    query_embedding_np = np.array([query_embedding]).astype("float32")
    D, I = st.session_state.faiss_index.search(query_embedding_np, k)

    bm25_scores = st.session_state.bm25.get_scores(query.split())

    hybrid_scores = {}
    for i in range(k):
        hybrid_scores[i] = bm25_scores[i] + D[0][i]

    sorted_indices = sorted(hybrid_scores, key=hybrid_scores.get, reverse=True)
    top_docs = [st.session_state.documents_list[i] for i in sorted_indices[:k]]

    return top_docs

if user_query:
    with st.spinner("Retrieving and processing query..."):
        retrieved_docs = hybrid_search(user_query, k=5)

        if retrieved_docs:
            context = "\n".join(retrieved_docs)

            input_data = {
                "context": context,
                "input": user_query
            }

            formatted_input = prompt.format(**input_data)

            response = st.session_state.llm.invoke(formatted_input)

            st.write("### Response:")
            st.write(response)

            json_response = {
                "query": user_query,
                "retrieved_documents": retrieved_docs,
                "llm_response": response.content,
            }

            with st.expander("JSON Response"):
                st.json(json_response)

            with st.expander("Document Similarity Search"):
                for doc in retrieved_docs:
                    st.write(doc)
                    st.write("---")
        else:
            st.warning("No relevant documents found.")
